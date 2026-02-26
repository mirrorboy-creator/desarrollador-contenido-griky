"""
File processor module for Desarrollador de Contenido Griky.
Handles reading PDF, DOCX, XLSX, and TXT files to extract course information.
"""
import io
from pathlib import Path
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error reading PDF: {str(e)}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text content from an XLSX file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Hoja: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Error reading XLSX: {str(e)}]"


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text content from a TXT file."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Error reading TXT: {str(e)}]"


def process_uploaded_file(filename: str, file_bytes: bytes) -> dict:
    """
    Process an uploaded file and return its content and metadata.

    Returns dict with keys: filename, extension, content, size_bytes
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".doc": extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".xls": extract_text_from_xlsx,
        ".txt": extract_text_from_txt,
    }

    extractor = extractors.get(ext, lambda b: b.decode("utf-8", errors="replace"))
    content = extractor(file_bytes)

    return {
        "filename": filename,
        "extension": ext,
        "content": content,
        "size_bytes": len(file_bytes),
        "char_count": len(content),
    }


def classify_document(filename: str, content: str) -> str:
    """
    Classify the document type based on filename and content keywords.
    Returns one of: syllabus, ebook_template, instructional_design, blueprint,
                    rubrics, competencies, activities, evaluations, other
    """
    name_lower = filename.lower()
    content_lower = content.lower()[:2000]

    classifications = {
        "syllabus": ["silabo", "syllabus", "programa", "plan de estudios"],
        "ebook_template": ["plantilla", "template", "ebook", "libro", "estructura"],
        "instructional_design": ["diseño instruccional", "instructional", "diseno instruccional"],
        "blueprint": ["blueprint", "mapa", "plano"],
        "rubrics": ["rubrica", "rúbrica", "rubric", "criterio", "evaluacion"],
        "competencies": ["competencia", "competencies", "resultado de aprendizaje", "learning outcome"],
        "activities": ["actividad", "activity", "tarea", "ejercicio"],
        "evaluations": ["evaluacion", "evaluation", "examen", "prueba", "assessment"],
    }

    for doc_type, keywords in classifications.items():
        for kw in keywords:
            if kw in name_lower or kw in content_lower:
                return doc_type

    return "other"


def build_course_context(processed_files: list[dict]) -> str:
    """
    Build a comprehensive course context string from all processed files.
    """
    sections = []
    sections.append("=== DOCUMENTOS DEL CURSO ===\n")

    for i, f in enumerate(processed_files, 1):
        doc_type = classify_document(f["filename"], f["content"])
        sections.append(
            f"--- Documento {i}: {f['filename']} [{doc_type.upper()}] ---\n"
            f"{f['content'][:8000]}\n"  # Limit per file to avoid token overflow
        )

    return "\n".join(sections)
