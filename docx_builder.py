"""
DOCX builder module for Desarrollador de Contenido Griky.
Generates formatted Word documents from generated academic content.
"""
import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GRIKY_PINK = RGBColor(0xF0, 0x3E, 0x5F)
GRIKY_DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRIKY_GRAY = RGBColor(0x4A, 0x4A, 0x6A)


def set_heading_style(paragraph, level: int, text: str):
    """Apply heading style to a paragraph."""
    paragraph.text = text
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = GRIKY_PINK
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = GRIKY_DARK
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = GRIKY_GRAY


def add_cover_page(doc: Document, course_name: str, subject_area: str):
    """Add a cover page to the document."""
    # Title section
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("DESARROLLADOR DE CONTENIDO ACADÉMICO")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = GRIKY_PINK

    doc.add_paragraph()

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("by Griky")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = GRIKY_GRAY
    subtitle_run.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    course_para = doc.add_paragraph()
    course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_run = course_para.add_run(course_name)
    course_run.font.size = Pt(20)
    course_run.font.bold = True
    course_run.font.color.rgb = GRIKY_DARK

    doc.add_paragraph()

    area_para = doc.add_paragraph()
    area_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    area_run = area_para.add_run(f"Área: {subject_area}")
    area_run.font.size = Pt(14)
    area_run.font.color.rgb = GRIKY_GRAY

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generado el {datetime.now().strftime('%d de %B de %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = GRIKY_GRAY

    doc.add_page_break()


def parse_content_sections(content: str) -> list[dict]:
    """
    Parse the generated content into sections (headings, paragraphs, image suggestions).
    Returns list of dicts: {type: 'heading1'|'heading2'|'heading3'|'paragraph'|'image'|'reference', text: str}
    """
    sections = []
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Image suggestion
        if stripped.startswith('[IMAGEN SUGERIDA:') or stripped.startswith('[IMAGEN SUGERIDA '):
            sections.append({"type": "image", "text": stripped})

        # Headings (markdown style)
        elif stripped.startswith('### '):
            sections.append({"type": "heading3", "text": stripped[4:]})
        elif stripped.startswith('## '):
            sections.append({"type": "heading2", "text": stripped[3:]})
        elif stripped.startswith('# '):
            sections.append({"type": "heading1", "text": stripped[2:]})

        # Bold headings (alternative format)
        elif re.match(r'^\*\*[A-Z]', stripped) and stripped.endswith('**'):
            text = stripped.strip('*')
            sections.append({"type": "heading2", "text": text})

        # References section
        elif stripped.lower() in ('referencias', 'referencias:', 'references', 'bibliografía'):
            sections.append({"type": "heading2", "text": stripped})

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('• '):
            sections.append({"type": "bullet", "text": stripped[2:]})
        elif re.match(r'^\d+\.\s', stripped):
            sections.append({"type": "numbered", "text": re.sub(r'^\d+\.\s', '', stripped)})

        # Regular paragraph
        else:
            sections.append({"type": "paragraph", "text": stripped})

    return sections


def add_image_suggestion_box(doc: Document, text: str):
    """Add a styled image suggestion box to the document."""
    # Add a bordered paragraph for image suggestion
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run("📷 ")
    run.font.size = Pt(11)

    desc_run = para.add_run(text.replace('[IMAGEN SUGERIDA:', '').replace(']', '').strip())
    desc_run.font.size = Pt(10)
    desc_run.font.italic = True
    desc_run.font.color.rgb = GRIKY_GRAY

    # Add light shading via XML manipulation
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4FF')
    pPr.append(shd)

    # Add border
    pBdr = OxmlElement('w:pBdr')
    for border_side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'F03E5F')
        pBdr.append(border)
    pPr.append(pBdr)


def clean_markdown_bold(text: str) -> list[tuple[str, bool]]:
    """
    Parse text with markdown bold (**text**) into list of (text, is_bold) tuples.
    """
    parts = []
    pattern = re.compile(r'\*\*(.*?)\*\*')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    return parts if parts else [(text, False)]


def add_paragraph_with_formatting(doc: Document, text: str, style: Optional[str] = None):
    """Add a paragraph with markdown bold formatting support."""
    para = doc.add_paragraph()
    if style:
        try:
            para.style = style
        except Exception:
            pass
    para.paragraph_format.space_after = Pt(6)

    parts = clean_markdown_bold(text)
    for part_text, is_bold in parts:
        run = para.add_run(part_text)
        run.font.size = Pt(11)
        if is_bold:
            run.font.bold = True
    return para


def build_docx(
    units_content: list[dict],
    course_name: str,
    subject_area: str,
) -> bytes:
    """
    Build a complete DOCX document from the generated content.

    Args:
        units_content: List of dicts with keys: unit_name, content, citations, citations_text
        course_name: Name of the course
        subject_area: Subject area/discipline

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Cover page
    add_cover_page(doc, course_name, subject_area)

    # Table of contents placeholder
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("TABLA DE CONTENIDO")
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True
    toc_run.font.color.rgb = GRIKY_PINK
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    for i, unit in enumerate(units_content, 1):
        toc_item = doc.add_paragraph(f"{i}. {unit['unit_name']}")
        toc_item.style = doc.styles['List Number'] if 'List Number' in [s.name for s in doc.styles] else doc.styles['Normal']
        toc_item.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # Generate content for each unit
    for unit_idx, unit in enumerate(units_content, 1):
        # Unit heading
        unit_heading = doc.add_paragraph()
        unit_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        unit_run = unit_heading.add_run(f"UNIDAD {unit_idx}: {unit['unit_name'].upper()}")
        unit_run.font.size = Pt(18)
        unit_run.font.bold = True
        unit_run.font.color.rgb = GRIKY_PINK

        # Decorative line
        line_para = doc.add_paragraph()
        line_run = line_para.add_run("─" * 70)
        line_run.font.color.rgb = GRIKY_PINK
        line_run.font.size = Pt(10)

        doc.add_paragraph()

        # Parse and add content sections
        sections = parse_content_sections(unit['content'])

        for section in sections:
            if section['type'] == 'heading1':
                h = doc.add_paragraph()
                run = h.add_run(section['text'])
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = GRIKY_PINK
                h.paragraph_format.space_before = Pt(12)

            elif section['type'] == 'heading2':
                h = doc.add_paragraph()
                run = h.add_run(section['text'])
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = GRIKY_DARK
                h.paragraph_format.space_before = Pt(10)

            elif section['type'] == 'heading3':
                h = doc.add_paragraph()
                run = h.add_run(section['text'])
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = GRIKY_GRAY
                h.paragraph_format.space_before = Pt(8)

            elif section['type'] == 'image':
                add_image_suggestion_box(doc, section['text'])

            elif section['type'] == 'bullet':
                bullet_para = doc.add_paragraph(style='List Bullet') if 'List Bullet' in [s.name for s in doc.styles] else doc.add_paragraph()
                parts = clean_markdown_bold(section['text'])
                for part_text, is_bold in parts:
                    run = bullet_para.add_run(part_text)
                    run.font.size = Pt(11)
                    if is_bold:
                        run.font.bold = True

            elif section['type'] == 'numbered':
                num_para = doc.add_paragraph(style='List Number') if 'List Number' in [s.name for s in doc.styles] else doc.add_paragraph()
                parts = clean_markdown_bold(section['text'])
                for part_text, is_bold in parts:
                    run = num_para.add_run(part_text)
                    run.font.size = Pt(11)
                    if is_bold:
                        run.font.bold = True

            else:  # paragraph
                add_paragraph_with_formatting(doc, section['text'])

        # Add page break between units (not after last one)
        if unit_idx < len(units_content):
            doc.add_page_break()

    # Final page: full references
    doc.add_page_break()
    ref_heading = doc.add_paragraph()
    ref_run = ref_heading.add_run("REFERENCIAS BIBLIOGRÁFICAS GENERALES")
    ref_run.font.size = Pt(16)
    ref_run.font.bold = True
    ref_run.font.color.rgb = GRIKY_PINK
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    all_citations_text = set()
    for unit in units_content:
        if unit.get('citations_text'):
            for line in unit['citations_text'].split('\n\n'):
                if line.strip():
                    all_citations_text.add(line.strip())

    if all_citations_text:
        for citation in sorted(all_citations_text):
            ref_para = doc.add_paragraph()
            ref_para.paragraph_format.left_indent = Inches(0.5)
            ref_para.paragraph_format.first_line_indent = Inches(-0.5)
            ref_run = ref_para.add_run(citation)
            ref_run.font.size = Pt(10)
    else:
        no_ref = doc.add_paragraph("No se obtuvieron referencias externas. Consulte las referencias indicadas en cada unidad.")
        no_ref.paragraph_format.left_indent = Inches(0.5)

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Generado por Desarrollador de Contenido Académico | by Griky")
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = GRIKY_GRAY

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
